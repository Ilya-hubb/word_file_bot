import os
from doAugust.src.states22 import WORDSTATE
from aiogram import Bot, Dispatcher, types, executor
from aiogram.dispatcher.filters.state import StatesGroup, State
from aiogram.types import ReplyKeyboardMarkup, KeyboardButton
from aiogram.contrib.fsm_storage.memory import MemoryStorage
from aiogram.dispatcher import FSMContext
import docx
from doAugust.generate_document_2 import generate_document
from TOKENAPI_word import TOKENAPI

TOKENAPI = TOKENAPI


storage: MemoryStorage = MemoryStorage()
bot = Bot(TOKENAPI)
dp = Dispatcher(bot,
                storage=storage)


def get_kb() -> ReplyKeyboardMarkup:
    kb = ReplyKeyboardMarkup(resize_keyboard=True)
    kb.add(KeyboardButton('/start'))

    return kb

def get_cancel_kb() -> ReplyKeyboardMarkup:
    kb = ReplyKeyboardMarkup(resize_keyboard=True)
    kb.add(KeyboardButton('/cancel'))

    return kb

class States(StatesGroup):
    text = State()

@dp.message_handler(commands=['cancel'], state='*')
async def cns_command(message: types.Message, state: FSMContext):
    if state is None:
        return

    await state.finish()
    await message.reply('Вы снова в главном меню!',
                        reply_markup=get_kb())

@dp.message_handler(commands=['start'])
async def start_func(message: types.Message) -> None:
    await message.reply(text='Привет! \nЯ готов к работе, для начала запонлни файл ниже и отправь его мне.')
    await bot.send_document(message.from_user.id, open(r'F:\file1.docx', 'rb'),
                            reply_markup=get_cancel_kb())

    await WORDSTATE.step1.set()

@dp.message_handler(content_types=['document'], state=WORDSTATE.step1)
async def load_text_from_doc(message: types.Message, state: FSMContext) -> None:
        if message.content_type == 'document':
            await message.document.download(destination_file=f'files/reciveddocc.docx')
            await message.answer('Ваш текст уже обрабатывается!')
            doc = docx.Document(f'files/reciveddocc.docx')
            text1 = (doc.paragraphs[0].text[10:])
            text2 = (doc.paragraphs[1].text[9:])
            text3 = (doc.paragraphs[2].text[9:])
            generate_document(text1, text2, text3)

            await bot.send_document(message.from_id, open(r'files/Rile.docx', 'rb'))
            os.remove('files/Rile.docx')
            os.remove('files/reciveddocc.docx')
        await state.finish()


if __name__ == '__main__':
    executor.start_polling(dp)

